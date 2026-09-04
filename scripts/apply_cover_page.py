"""Insert a CDU cover page (logo + centred fields) into the report .docx."""

from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

ROOT = Path(__file__).resolve().parents[1]
LOGO = ROOT / "docs" / "assets" / "cdu-logo.png"

COVER_LINES = [
    ("Software Engineering: Process and Tools", 12, False, 18),
    ("Unit Code: PRT582", 12, False, 10),
    ("Software Unit Testing Report", 14, True, 16),
    ("Name: Saugat Shrestha", 12, False, 8),
    ("Student ID: S403036", 12, False, 4),
    ("Campus: Sydney", 12, False, 4),
    ("Submission Date: 2 September 2026", 12, False, 4),
    ("Submitted To: Abdullah Al-Amoodi", 12, False, 4),
    (
        "Application built: Python Static Code Analyzer (\"PyScan\")",
        11,
        False,
        10,
    ),
    (
        "GitHub repository: https://github.com/Saugat-Shrestha/python_static_code_analyzer",
        11,
        False,
        4,
    ),
]

COVER_PREFIXES = (
    "Software Engineering: Process and Tools",
    "Unit Code: PRT582",
    "Software Unit Testing Report",
    "Name:",
    "Student ID:",
    "Campus:",
    "Submission Date:",
    "Submitted To:",
    "Application built:",
    "GitHub repository:",
)


def _set_paragraph_spacing(paragraph, before_pt: float, after_pt: float) -> None:
    paragraph.paragraph_format.space_before = Pt(before_pt)
    paragraph.paragraph_format.space_after = Pt(after_pt)
    paragraph.paragraph_format.line_spacing = 1.0


def _add_page_break(paragraph) -> None:
    run = paragraph.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._r.append(br)


def _remove_existing_cover(document: Document) -> None:
    while document.paragraphs:
        text = document.paragraphs[0].text.strip()
        if not text:
            p = document.paragraphs[0]._element
            p.getparent().remove(p)
            continue
        if any(text.startswith(prefix) for prefix in COVER_PREFIXES):
            p = document.paragraphs[0]._element
            p.getparent().remove(p)
            continue
        break


def apply_cover(docx_path: Path) -> None:
    if not LOGO.exists():
        raise FileNotFoundError(f"Logo not found: {LOGO}")

    document = Document(str(docx_path))
    _remove_existing_cover(document)

    body = document.element.body
    first = body[0]

    logo_para = document.add_paragraph()
    logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_spacing(logo_para, 0, 18)
    logo_para.add_run().add_picture(str(LOGO), width=Cm(4.5))
    body.insert(body.index(first), logo_para._element)

    for text, size, bold, after in COVER_LINES:
        para = document.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        before = 14 if text.startswith("Software Engineering") else 0
        _set_paragraph_spacing(para, before, after)
        run = para.add_run(text)
        run.font.size = Pt(size)
        run.bold = bold
        body.insert(body.index(first), para._element)

    break_para = document.add_paragraph()
    _add_page_break(break_para)
    body.insert(body.index(first), break_para._element)

    document.save(str(docx_path))


def main() -> None:
    docx = Path(sys.argv[1]) if len(sys.argv) > 1 else ROOT / (
        "PRT582 Software Unit Testing Report - Saugat Shrestha S403036.docx"
    )
    apply_cover(docx)
    print(f"Updated cover page in {docx}")


if __name__ == "__main__":
    main()
